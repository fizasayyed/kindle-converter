from docx import Document
import pytesseract
from PIL import Image
import io

# Load the DOCX file
input_doc = "/input/input.docx"  # Change to your filename
output_doc = "converted_op.docx"

doc = Document(input_doc)
new_doc = Document()

# Process each element while preserving structure
for para in doc.paragraphs:
    new_doc.add_paragraph(para.text)  # Preserve existing text

for table in doc.tables:
    new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            new_table.cell(i, j).text = cell.text  # Copy table text

# Extract and process images
for rel in doc.part.rels:
    if "image" in doc.part.rels[rel].target_ref:
        image = doc.part.rels[rel].target_part.blob
        img = Image.open(io.BytesIO(image))

        # Extract text using OCR
        extracted_text = pytesseract.image_to_string(img)

        # Add image back to the document
        new_doc.add_picture(io.BytesIO(image))

        # Insert OCR text below the image
        new_doc.add_paragraph(extracted_text)

# Save the processed DOCX file
new_doc.save(output_doc)
print(f"OCR Completed. File saved as {output_doc}")