import os
import zipfile
import tempfile
import subprocess
from docx import Document
from PIL import Image
import pytesseract
import sys

# Configure Tesseract path (uncomment and modify if needed)
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def has_text(docx_path, threshold=100):
    """Check if document contains sufficient text"""
    doc = Document(docx_path)
    text = '\n'.join([para.text for para in doc.paragraphs])
    return len(text.strip()) >= threshold

def extract_images(docx_path, output_folder):
    """Extract images from DOCX in correct order"""
    with zipfile.ZipFile(docx_path, 'r') as docx_zip:
        image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
        image_files.sort()

        for idx, img_path in enumerate(image_files, 1):
            img_data = docx_zip.read(img_path)
            output_path = os.path.join(output_folder, f"page_{idx}{os.path.splitext(img_path)[1]}")
            with open(output_path, 'wb') as f:
                f.write(img_data)
    return len(image_files)

def ocr_images(image_folder, output_docx):
    """Perform OCR on images and create new DOCX"""
    doc = Document()
    image_files = sorted([f for f in os.listdir(image_folder) if f.startswith('page_')],
                         key=lambda x: int(x.split('_')[1].split('.')[0]))

    for img_file in image_files:
        img_path = os.path.join(image_folder, img_file)
        try:
            text = pytesseract.image_to_string(Image.open(img_path))
            doc.add_paragraph(text)
            doc.add_page_break()
        except Exception as e:
            print(f"Error processing {img_file}: {str(e)}")

    doc.save(output_docx)

def convert_to_mobi(input_path, output_path):
    """Convert document to MOBI using Calibre"""
    try:
        subprocess.run(['ebook-convert', input_path, output_path], check=True)
    except subprocess.CalledProcessError as e:
        print(f"Error converting to MOBI: {e}")
        sys.exit(1)
    except FileNotFoundError:
        print("Calibre not found. Install Calibre and ensure 'ebook-convert' is in PATH")
        sys.exit(1)

def main(docx_path, mobi_path):
    if has_text(docx_path):
        print("Document contains text. Converting directly to MOBI...")
        convert_to_mobi(docx_path, mobi_path)
    else:
        print("Document appears to be image-based. Performing OCR...")
        with tempfile.TemporaryDirectory() as temp_dir:
            image_dir = os.path.join(temp_dir, 'images')
            os.makedirs(image_dir, exist_ok=True)

            num_images = extract_images(docx_path, image_dir)
            print(f"Extracted {num_images} images for OCR processing.")

            ocr_docx = os.path.join(temp_dir, 'ocr_output.docx')
            ocr_images(image_dir, ocr_docx)

            print("OCR completed. Converting to MOBI...")
            convert_to_mobi(ocr_docx, mobi_path)

    print("Conversion to MOBI complete. Output file:", mobi_path)

if __name__ == "__main__":
    input_docx = "/input/input.docx"  # Replace with your DOCX path
    output_mobi = "output.mobi"  # Replace with desired output path

    main(input_docx, output_mobi)